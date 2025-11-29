"""Document processing utilities for the AI-Enhanced Interactive Book Agent.

This module provides utilities for processing various document formats (PDF, DOCX, EPUB, etc.)
for use in the RAG system, including text extraction, cleaning, and preparation for chunking.
"""
import asyncio
import os
import tempfile
from typing import List, Tuple, Dict, Any, Optional
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import chardet
from langchain.text_splitter import RecursiveCharacterTextSplitter


class DocumentProcessor:
    """Main class for processing documents of various formats."""

    def __init__(self):
        """Initialize the document processor."""
        self.supported_formats = ['.pdf', '.docx', '.epub', '.txt']
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            length_function=len,
        )

    async def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a document and extract its content.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary containing extracted content and metadata
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")

        file_extension = file_path.suffix.lower()
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported: {self.supported_formats}")

        # Extract content based on file type
        if file_extension == '.pdf':
            content, metadata = await self._extract_pdf_content(file_path)
        elif file_extension == '.docx':
            content, metadata = await self._extract_docx_content(file_path)
        elif file_extension == '.epub':
            content, metadata = await self._extract_epub_content(file_path)
        elif file_extension == '.txt':
            content, metadata = await self._extract_txt_content(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

        # Clean and preprocess the content
        cleaned_content = self._clean_text(content)

        # Extract additional metadata
        file_stats = file_path.stat()
        additional_metadata = {
            "file_name": file_path.name,
            "file_path": str(file_path),
            "file_size": file_stats.st_size,
            "file_extension": file_extension,
            "created_at": file_stats.st_ctime,
            "modified_at": file_stats.st_mtime
        }

        # Combine metadata
        full_metadata = {**metadata, **additional_metadata}

        return {
            "content": cleaned_content,
            "metadata": full_metadata,
            "chunks": await self._chunk_content(cleaned_content, full_metadata)
        }

    async def _extract_pdf_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from a PDF file.

        Args:
            file_path: Path to the PDF file

        Returns:
            Tuple of (content, metadata)
        """
        content = ""
        metadata = {}

        try:
            # Using pdfplumber for better text extraction
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n"

            # Extract metadata using PyPDF2
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                pdf_metadata = pdf_reader.metadata

                if pdf_metadata:
                    metadata = {
                        "title": pdf_metadata.get('/Title', ''),
                        "author": pdf_metadata.get('/Author', ''),
                        "subject": pdf_metadata.get('/Subject', ''),
                        "creator": pdf_metadata.get('/Creator', ''),
                        "producer": pdf_metadata.get('/Producer', ''),
                        "creation_date": str(pdf_metadata.get('/CreationDate', '')),
                        "modification_date": str(pdf_metadata.get('/ModDate', '')),
                        "pages": len(pdf_reader.pages)
                    }

        except Exception as e:
            print(f"Error processing PDF: {str(e)}")
            # Fallback using PyPDF2 only
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"

        return content, metadata

    async def _extract_docx_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (content, metadata)
        """
        content = ""
        metadata = {}

        try:
            doc = DocxDocument(file_path)

            # Extract content
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"

            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "comments": core_props.comments or "",
                "keywords": core_props.keywords or "",
                "last_modified_by": core_props.last_modified_by or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }

        except Exception as e:
            print(f"Error processing DOCX: {str(e)}")
            raise e

        return content, metadata

    async def _extract_epub_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from an EPUB file.

        Args:
            file_path: Path to the EPUB file

        Returns:
            Tuple of (content, metadata)
        """
        content = ""
        metadata = {}

        try:
            book = epub.read_epub(file_path)

            # Extract metadata
            metadata = {
                "title": list(book.get_metadata('DC', 'title'))[0][0] if book.get_metadata('DC', 'title') else "",
                "creator": list(book.get_metadata('DC', 'creator'))[0][0] if book.get_metadata('DC', 'creator') else "",
                "subject": list(book.get_metadata('DC', 'subject'))[0][0] if book.get_metadata('DC', 'subject') else "",
                "description": list(book.get_metadata('DC', 'description'))[0][0] if book.get_metadata('DC', 'description') else "",
                "publisher": list(book.get_metadata('DC', 'publisher'))[0][0] if book.get_metadata('DC', 'publisher') else "",
                "contributor": list(book.get_metadata('DC', 'contributor'))[0][0] if book.get_metadata('DC', 'contributor') else "",
                "date": list(book.get_metadata('DC', 'date'))[0][0] if book.get_metadata('DC', 'date') else "",
                "type": list(book.get_metadata('DC', 'type'))[0][0] if book.get_metadata('DC', 'type') else "",
                "format": list(book.get_metadata('DC', 'format'))[0][0] if book.get_metadata('DC', 'format') else "",
                "identifier": list(book.get_metadata('DC', 'identifier'))[0][0] if book.get_metadata('DC', 'identifier') else "",
                "source": list(book.get_metadata('DC', 'source'))[0][0] if book.get_metadata('DC', 'source') else "",
                "language": list(book.get_metadata('DC', 'language'))[0][0] if book.get_metadata('DC', 'language') else "",
                "relation": list(book.get_metadata('DC', 'relation'))[0][0] if book.get_metadata('DC', 'relation') else "",
                "coverage": list(book.get_metadata('DC', 'coverage'))[0][0] if book.get_metadata('DC', 'coverage') else "",
                "rights": list(book.get_metadata('DC', 'rights'))[0][0] if book.get_metadata('DC', 'rights') else "",
            }

            # Extract content from chapters
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    content += soup.get_text() + "\n"

        except Exception as e:
            print(f"Error processing EPUB: {str(e)}")
            raise e

        return content, metadata

    async def _extract_txt_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from a TXT file.

        Args:
            file_path: Path to the TXT file

        Returns:
            Tuple of (content, metadata)
        """
        content = ""
        metadata = {"encoding": "unknown"}

        try:
            # Try common encodings first
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        metadata["encoding"] = encoding
                        break
                except UnicodeDecodeError:
                    continue
            else:
                # If common encodings fail, try to detect encoding
                with open(file_path, 'rb') as file:
                    raw_data = file.read()
                    encoding_result = chardet.detect(raw_data)
                    detected_encoding = encoding_result['encoding']
                    
                    if detected_encoding:
                        content = raw_data.decode(detected_encoding)
                        metadata["encoding"] = detected_encoding
                    else:
                        raise UnicodeDecodeError("Could not determine file encoding")

        except Exception as e:
            print(f"Error processing TXT: {str(e)}")
            raise e

        return content, metadata

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text.

        Args:
            text: Raw text to clean

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Normalize whitespace
        text = ' '.join(text.split())

        # Remove extra newlines but preserve paragraph structure
        import re
        text = re.sub(r'\n\s*\n', '\n\n', text)

        # Remove control characters except newlines and tabs
        text = ''.join(char for char in text if ord(char) < 32 and char not in '\n\t' or ord(char) >= 32)

        return text.strip()

    async def _chunk_content(self, content: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split content into chunks for RAG processing.

        Args:
            content: Content to chunk
            metadata: Metadata associated with the content

        Returns:
            List of chunk dictionaries
        """
        chunks = self.text_splitter.split_text(content)
        chunked_data = []

        for idx, chunk in enumerate(chunks):
            chunk_data = {
                "chunk_id": f"{metadata.get('file_name', 'unknown')}_{idx}",
                "content": chunk,
                "metadata": {
                    **metadata,
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                    "chunk_size": len(chunk),
                }
            }
            chunked_data.append(chunk_data)

        return chunked_data

    async def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic information about a document without processing its full content.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary with document information
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")

        file_extension = file_path.suffix.lower()
        file_stats = file_path.stat()

        info = {
            "file_name": file_path.name,
            "file_path": str(file_path),
            "file_size": file_stats.st_size,
            "file_extension": file_extension,
            "created_at": file_stats.st_ctime,
            "modified_at": file_stats.st_mtime,
            "is_supported": file_extension in self.supported_formats
        }

        # Try to get basic metadata without full processing
        if file_extension == '.pdf':
            try:
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    info["pages"] = len(pdf_reader.pages)
                    pdf_metadata = pdf_reader.metadata
                    if pdf_metadata:
                        info["title"] = pdf_metadata.get('/Title', '')
                        info["author"] = pdf_metadata.get('/Author', '')
            except Exception:
                info["pages"] = "unknown"
        elif file_extension == '.docx':
            try:
                doc = DocxDocument(file_path)
                info["pages"] = doc.core_properties.pages if doc.core_properties.pages else "unknown"
                info["title"] = doc.core_properties.title or ""
                info["author"] = doc.core_properties.author or ""
            except Exception:
                info["pages"] = "unknown"
        elif file_extension == '.epub':
            try:
                book = epub.read_epub(file_path)
                info["title"] = list(book.get_metadata('DC', 'title'))[0][0] if book.get_metadata('DC', 'title') else ""
            except Exception:
                info["title"] = "unknown"
        elif file_extension == '.txt':
            info["line_count"] = sum(1 for line in open(file_path, 'r', encoding='utf-8'))

        return info


class BatchDocumentProcessor:
    """Class for processing multiple documents in batch."""

    def __init__(self):
        """Initialize the batch document processor."""
        self.processor = DocumentProcessor()

    async def process_documents(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """Process multiple documents asynchronously.

        Args:
            file_paths: List of paths to document files

        Returns:
            List of processed document results
        """
        tasks = [self.processor.process_document(file_path) for file_path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                print(f"Error processing {file_paths[i]}: {str(result)}")
                processed_results.append({
                    "file_path": file_paths[i],
                    "error": str(result),
                    "success": False
                })
            else:
                result["success"] = True
                processed_results.append(result)

        return processed_results


# Global instance of the document processor
document_processor = DocumentProcessor()
batch_document_processor = BatchDocumentProcessor()